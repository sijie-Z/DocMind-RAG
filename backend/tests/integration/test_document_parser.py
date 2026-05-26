"""DocumentParser 纯函数测试（不依赖外部服务）。"""
import os
import tempfile

import pytest

from app.services.document_parser import DocumentParser


@pytest.fixture
def parser():
    return DocumentParser()


class TestParseTxt:
    @pytest.mark.asyncio
    async def test_parse_plain_text(self, parser: DocumentParser):
        with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False, encoding="utf-8") as f:
            f.write("这是一段测试文本。\n第二行内容。")
            f.flush()
            path = f.name
        try:
            result = await parser._parse_txt(path)
            assert "测试文本" in result["content"]
            assert "encoding" in result["metadata"]
            assert result["metadata"]["encoding"] == "utf-8"
        finally:
            os.unlink(path)

    @pytest.mark.asyncio
    async def test_parse_markdown(self, parser: DocumentParser):
        with tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False, encoding="utf-8") as f:
            f.write("# 标题\n\n正文内容\n\n## 子标题\n\n更多内容。")
            f.flush()
            path = f.name
        try:
            result = await parser._parse_txt(path)
            assert "标题" in result["content"]
            assert result["metadata"]["line_count"] >= 4
        finally:
            os.unlink(path)

    @pytest.mark.asyncio
    async def test_parse_empty_txt(self, parser: DocumentParser):
        with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False, encoding="utf-8") as f:
            f.write("")
            f.flush()
            path = f.name
        try:
            result = await parser._parse_txt(path)
            assert result["content"] == ""
        finally:
            os.unlink(path)

    @pytest.mark.asyncio
    async def test_parse_txt_metadata(self, parser: DocumentParser):
        content = "第一行\n第二行\n第三行"
        with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False, encoding="utf-8") as f:
            f.write(content)
            f.flush()
            path = f.name
        try:
            result = await parser._parse_txt(path)
            assert result["metadata"]["char_count"] == len(content)
            assert result["metadata"]["line_count"] == 3
        finally:
            os.unlink(path)


class TestParseDocx:
    @pytest.mark.asyncio
    async def test_parse_docx(self, parser: DocumentParser):
        """Create a minimal .docx and parse it."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            pytest.skip("python-docx not installed")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name

        doc = DocxDocument()
        doc.add_paragraph("第一段内容。")
        doc.add_paragraph("第二段内容。")
        doc.save(path)

        try:
            result = await parser._parse_docx(path)
            assert "第一段内容" in result["content"]
            assert "第二段内容" in result["content"]
            assert "paragraphs" in result["metadata"]
        finally:
            os.unlink(path)


class TestSupportedContent:
    def test_supported_extensions_include_common_types(self, parser: DocumentParser):
        assert ".pdf" in parser.supported_extensions
        assert ".docx" in parser.supported_extensions
        assert ".txt" in parser.supported_extensions
        assert ".md" in parser.supported_extensions
        assert ".xlsx" in parser.supported_extensions
