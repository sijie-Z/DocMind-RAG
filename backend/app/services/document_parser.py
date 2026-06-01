"""
文档解析服务 - 支持PDF、Word、Excel、TXT等多种格式
"""
import asyncio
import contextlib
import logging
import os
import uuid
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET

import docx2txt
import pandas as pd
import pdfplumber
from docx import Document as DocxDocument
from openpyxl import load_workbook
from sqlalchemy import select
from sqlalchemy.orm import selectinload

from app.core.config import settings
from app.core.database import AsyncSessionLocal
from app.models.document import Document, DocumentChunk, DocumentStatus

logger = logging.getLogger(__name__)


class DocumentParser:
    """文档解析器 - 支持多种文件格式"""

    def __init__(self):
        self.supported_extensions = {
            '.pdf': self._parse_pdf,
            '.docx': self._parse_docx,
            '.xlsx': self._parse_excel,
            '.xls': self._parse_excel,
            '.txt': self._parse_txt,
            '.md': self._parse_txt,
        }

    async def parse_document(self, file_path: str, organization_id: str) -> dict[str, Any]:
        """
        解析文档
        
        Args:
            file_path: 文件路径 (可以是本地路径，也可以是 MinIO 对象名称或 URL)
            organization_id: 组织ID
            
        Returns:
            解析结果，包含内容 and 元数据
        """
        import tempfile

        from app.core.minio_client import minio_client

        temp_file_path = None
        try:
            # 1. 判断并获取文件本地路径
            if file_path.startswith('http://') or file_path.startswith('https://') or '/' in file_path:
                # 可能是 URL 或 MinIO 路径，需要先下载到本地临时文件
                logger.info(f"正在从存储下载文件以供解析: {file_path}")

                # 提取对象名 (如果是 URL)
                object_name = file_path
                if "://" in file_path:
                    # 假设 URL 格式为 http://endpoint/bucket/object_name
                    parts = file_path.split("/")
                    object_name = "/".join(parts[4:]) if len(parts) > 4 else parts[-1]

                # 创建临时文件
                ext = Path(file_path).suffix.lower() or ".tmp"
                fd, temp_file_path = tempfile.mkstemp(suffix=ext)
                os.close(fd)

                # 下载文件
                try:
                    await asyncio.to_thread(minio_client.fget_object, settings.MINIO_BUCKET_NAME, object_name, temp_file_path)
                    current_file_path = temp_file_path
                except Exception as e:
                    logger.error(f"从 MinIO 下载文件失败: {str(e)}")
                    # 兜底：如果 fget_object 失败，尝试直接作为本地路径（虽然大概率也会失败）
                    current_file_path = file_path
            else:
                current_file_path = file_path

            path = Path(current_file_path)
            if not path.exists():
                raise FileNotFoundError(f"文件不存在: {current_file_path}")

            # 获取文件扩展名
            extension = path.suffix.lower()

            # 检查是否支持该格式
            if extension not in self.supported_extensions:
                raise ValueError(f"不支持的文件格式: {extension}")

            # 调用对应的解析方法
            parse_method = self.supported_extensions[extension]
            result = await parse_method(current_file_path)

            # 添加通用元数据
            result["file_name"] = path.name
            result["file_size"] = path.stat().st_size
            result["file_extension"] = extension
            result["organization_id"] = organization_id

            logger.info(f"文档解析成功: {file_path}")
            return result

        except Exception as e:
            logger.error(f"文档解析失败: {file_path}, 错误: {str(e)}")
            raise
        finally:
            # 清理临时文件
            if temp_file_path and os.path.exists(temp_file_path):
                with contextlib.suppress(OSError):
                    os.remove(temp_file_path)

    async def parse_local_file(self, file_path: str, organization_id: str = "0") -> dict[str, Any]:
        """解析本地文件，不经过 MinIO 下载逻辑（用于聊天即时解析）"""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        extension = path.suffix.lower()
        if extension not in self.supported_extensions:
            raise ValueError(f"不支持的文件格式: {extension}")

        parse_method = self.supported_extensions[extension]
        result = await parse_method(str(path))
        result["file_name"] = path.name
        result["file_size"] = path.stat().st_size
        result["file_extension"] = extension
        result["organization_id"] = organization_id

        return result

    async def _parse_pdf(self, file_path: str) -> dict[str, Any]:
        """
        高级 PDF 解析 - 支持文本提取、表格识别与 OCR 兜底
        """
        try:
            text_content = ""
            metadata: dict[str, Any] = {"pages":[], "tables_count": 0, "ocr_used": False}

            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # 1. 提取页面文本
                    extracted_text = page.extract_text()
                    page_text: str = str(extracted_text) if extracted_text else ""

                    # 2. 如果文本太少，尝试 OCR (扫描件识别)
                    if len(page_text.strip()) < 50:
                        try:
                            import pytesseract
                            from pdf2image import convert_from_path

                            logger.info(f"📄 页面 {page_num} 文本过少，启动 OCR...")
                            # 仅转换当前页
                            images = convert_from_path(file_path, first_page=page_num, last_page=page_num, dpi=200)
                            if images:
                                ocr_text = pytesseract.image_to_string(images[0], lang=getattr(settings, "OCR_LANGUAGE", "chi_sim+eng"))
                                if ocr_text.strip():
                                    page_text = ocr_text
                                    metadata["ocr_used"] = True
                                    logger.info(f"✅ OCR 识别成功，长度: {len(page_text)}")
                        except Exception as ocr_err:
                            logger.warning(f"OCR 失败: {ocr_err}")

                    # 3. 提取表格并转换为 Markdown
                    tables = page.extract_tables()
                    table_mds =[]
                    for table in tables:
                        if table:
                            df = pd.DataFrame(table)
                            df = df.dropna(how='all').fillna('')
                            if not df.empty:
                                markdown_str = df.to_markdown(index=False)
                                if markdown_str is not None:
                                    table_md = "\n" + str(markdown_str) + "\n"
                                    table_mds.append(table_md)
                                    metadata["tables_count"] += 1

                    # 4. 合并文本和表格
                    combined_text: str = page_text
                    for t_md in table_mds:
                        combined_text += t_md

                    text_content += combined_text + "\n\n"

                    metadata["pages"].append({
                        "page_number": page_num,
                        "text_length": len(page_text),
                        "tables_found": len(table_mds)
                    })

            return {
                "content": text_content.strip(),
                "content_length": len(text_content),
                "metadata": metadata,
                "chunks": await self._create_contextual_chunks(text_content, metadata)
            }
        except Exception as e:
            logger.error(f"PDF 解析失败: {str(e)}")
            raise Exception(f"PDF 解析失败: {str(e)}")

    async def _parse_docx(self, file_path: str) -> dict[str, Any]:
        """解析Word文档"""
        text_content = ""
        metadata: dict[str, Any] = {"paragraphs": [], "tables":[]}
        try:
            doc = DocxDocument(file_path)

            # 提取段落文本
            for para_num, paragraph in enumerate(doc.paragraphs, 1):
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n\n"

                    # 记录段落信息
                    para_info = {
                        "paragraph_number": para_num,
                        "text_length": len(paragraph.text),
                        "style": paragraph.style.name if paragraph.style else "Normal"
                    }
                    metadata["paragraphs"].append(para_info)

            # 提取表格内容
            for table_num, table in enumerate(doc.tables, 1):
                table_text = ""
                for row in table.rows:
                    row_text =[]
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    table_text += " | ".join(row_text) + "\n"

                if table_text.strip():
                    text_content += f"\n【表格 {table_num}】\n{table_text}\n"

                    # 记录表格信息
                    table_info = {
                        "table_number": table_num,
                        "rows": len(table.rows),
                        "columns": len(table.columns) if table.rows else 0,
                        "text_length": len(table_text)
                    }
                    metadata["tables"].append(table_info)

            # 提取文档属性
            if doc.core_properties:
                metadata.update({
                    "title": doc.core_properties.title or "",
                    "author": doc.core_properties.author or "",
                    "subject": doc.core_properties.subject or "",
                    "created": str(doc.core_properties.created) if doc.core_properties.created else "",
                    "modified": str(doc.core_properties.modified) if doc.core_properties.modified else ""
                })
        except Exception as e:
            logger.warning(f"python-docx 解析失败，尝试回退方案: {str(e)}")

        if len(text_content.strip()) < 20:
            try:
                fallback = docx2txt.process(file_path) or ""
                if fallback.strip():
                    text_content = fallback.strip()
                    metadata["fallback_parser"] = "docx2txt"
            except Exception as e:
                logger.warning(f"docx2txt 解析失败: {str(e)}")

        if len(text_content.strip()) < 20:
            xml_text = self._extract_docx_text_from_xml(file_path)
            if xml_text.strip():
                text_content = xml_text.strip()
                metadata["fallback_parser"] = "zip_xml"

        if not text_content.strip():
            raise Exception("DOCX解析失败：未提取到有效文本内容")

        return {
            "content": text_content.strip(),
            "content_length": len(text_content),
            "metadata": metadata,
            "chunks": await self._create_contextual_chunks(text_content, metadata)
        }

    def _extract_docx_text_from_xml(self, file_path: str) -> str:
        try:
            if not zipfile.is_zipfile(file_path):
                return ""
            with zipfile.ZipFile(file_path, "r") as zf:
                if "word/document.xml" not in zf.namelist():
                    return ""
                xml_bytes = zf.read("word/document.xml")
            root = ET.fromstring(xml_bytes)
            namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            texts = [node.text for node in root.findall(".//w:t", namespace) if node.text]
            return "\n".join([t.strip() for t in texts if t and t.strip()])
        except Exception as e:
            logger.warning(f"zip_xml 解析失败: {str(e)}")
            return ""

    async def _parse_excel(self, file_path: str) -> dict[str, Any]:
        """解析Excel文件"""
        try:
            text_content = ""
            metadata: dict[str, Any] = {"sheets":[]}

            # 加载工作簿
            workbook = load_workbook(file_path, read_only=True)

            # 遍历每个工作表
            for sheet_name in workbook.sheetnames:
                worksheet = workbook[sheet_name]
                sheet_text = f"【工作表: {sheet_name}】\n\n"

                # 提取表格数据
                for _row_num, row in enumerate(worksheet.iter_rows(values_only=True), 1):
                    row_data =[]
                    for cell_value in row:
                        if cell_value is not None:
                            row_data.append(str(cell_value))

                    if row_data:
                        sheet_text += " | ".join(row_data) + "\n"

                if sheet_text.strip():
                    text_content += sheet_text + "\n"

                    # 记录工作表信息
                    sheet_info = {
                        "sheet_name": sheet_name,
                        "max_row": worksheet.max_row,
                        "max_column": worksheet.max_column,
                        "text_length": len(sheet_text)
                    }
                    metadata["sheets"].append(sheet_info)

            workbook.close()

            return {
                "content": text_content.strip(),
                "content_length": len(text_content),
                "metadata": metadata,
                "chunks": await self._create_contextual_chunks(text_content, metadata)
            }

        except Exception as e:
            raise Exception(f"Excel文件解析失败: {str(e)}")

    async def _parse_txt(self, file_path: str) -> dict[str, Any]:
        """解析文本文件"""
        try:
            # 检测文件编码
            encodings =['utf-8', 'gbk', 'gb2312', 'utf-16']
            text_content = None
            encoding_used = "utf-8"

            for encoding in encodings:
                try:
                    with open(file_path, encoding=encoding) as file:
                        text_content = file.read()
                        encoding_used = encoding
                    break
                except UnicodeDecodeError:
                    continue

            if text_content is None:
                raise Exception("无法识别的文件编码")

            metadata: dict[str, Any] = {
                "encoding": encoding_used,
                "line_count": len(text_content.split('\n')),
                "char_count": len(text_content)
            }

            return {
                "content": text_content.strip(),
                "content_length": len(text_content),
                "metadata": metadata,
                "chunks": await self._create_contextual_chunks(text_content, metadata)
            }

        except Exception as e:
            raise Exception(f"文本文件解析失败: {str(e)}")

    async def _generate_summary(self, text: str) -> str:
        """生成文档的简短摘要，用于上下文增强检索。"""
        if not text or len(text) < 1000:
            return ""

        try:
            from app.services.rag_service import rag_service
            if not rag_service.openai_client:
                return ""

            prompt = (
                "请为以下文档内容生成一段非常简短的背景介绍（50-100字）。"
                "这段文字将作为上下文附加到文档的每一个片段中，以帮助检索模型理解其背景。\n\n"
                f"文档内容: {text[:3000]}..."
            )

            model = settings.LOCAL_LLM_MODEL if settings.ENABLE_LOCAL_LLM else settings.DEEPSEEK_MODEL
            resp = await rag_service.openai_client.chat.completions.create(
                model=model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3,
                max_tokens=200
            )
            return (resp.choices[0].message.content or "").strip()
        except Exception as e:
            logger.warning(f"Summary generation failed: {e}")
            return ""

    async def _create_contextual_chunks(self, text_content: str, metadata: dict[str, Any]) -> list[dict[str, Any]]:
        """实现上下文增强分块 (Contextual Retrieval)，融合条款感知切分。"""
        # 1. 生成文档级摘要
        summary = await self._generate_summary(text_content)

        # 2. 执行语义分块（返回 list[dict]，每个含 text + metadata.section_title）
        from app.services.semantic_chunker import semantic_chunker
        semantic_chunks = await semantic_chunker.split_text(text_content)

        chunks = []
        for i, ch in enumerate(semantic_chunks):
            chunk_text = ch["text"]
            chunk_meta = ch.get("metadata", {})
            section_title = chunk_meta.get("section_title")

            # 3. 增强文本：[背景摘要] + 原文本
            if summary:
                enhanced_text = f"【文档背景：{summary}】\n\n{chunk_text}"
            else:
                enhanced_text = chunk_text

            chunks.append({
                "chunk_index": i,
                "chunk_text": enhanced_text,
                "original_text": chunk_text,
                "chunk_length": len(enhanced_text),
                "metadata": {
                    "type": "contextual_semantic_split",
                    "has_summary": bool(summary),
                    "section_title": section_title or chunk_meta.get("section_title"),
                }
            })
        return chunks

    async def process_document(self, document_id: str) -> bool:
        """
        处理文档 - 解析内容并创建分块
        
        Args:
            document_id: 文档ID
            
        Returns:
            处理成功返回True，失败返回False
        """
        try:
            async with AsyncSessionLocal() as session:
                # 获取文档信息
                result = await session.execute(
                    select(Document)
                    .where(Document.id == document_id)
                    .options(selectinload(Document.chunks))
                )
                document = result.scalar_one_or_none()

                if not document:
                    logger.error(f"文档不存在: {document_id}")
                    return False

                # 更新文档状态为处理中（忽略 SQLAlchemy Column 类型提示报警）
                document.status = DocumentStatus.PARSING
                await session.commit()

                try:
                    # 解析文档
                    parse_result = await self.parse_document(
                        document.file_path,
                        document.organization_id
                    )

                    # 更新文档信息
                    document.status = DocumentStatus.PARSED  # type: ignore
                    document.content_length = parse_result["content_length"]  # type: ignore
                    document.chunk_count = len(parse_result["chunks"])  # type: ignore
                    document.parsed_at = datetime.now()  # type: ignore

                    # 保存文档分块
                    for chunk_data in parse_result["chunks"]:
                        chunk = DocumentChunk(
                            id=str(uuid.uuid4()),
                            document_id=document_id,
                            chunk_index=chunk_data["chunk_index"],
                            chunk_text=chunk_data["chunk_text"],
                            chunk_length=chunk_data["chunk_length"],
                            start_pos=chunk_data.get("start_pos"),
                            end_pos=chunk_data.get("end_pos"),
                            page_number=chunk_data["metadata"].get("page_number"),
                            section_title=chunk_data["metadata"].get("section_title"),
                            meta_data=chunk_data["metadata"]
                        )
                        session.add(chunk)

                    await session.commit()
                    logger.info(f"文档处理完成: {document_id}")
                    return True

                except Exception as e:
                    await session.rollback()

                    # 更新文档状态为失败
                    if document:
                        document.status = DocumentStatus.FAILED  # type: ignore
                        document.parse_error = str(e)  # type: ignore
                        await session.commit()

                    logger.error(f"文档处理失败: {document_id}, 错误: {str(e)}")
                    return False

        except Exception as e:
            logger.error(f"处理文档时发生异常: {document_id}, 错误: {str(e)}")
            return False

    async def get_document_chunks(self, document_id: str) -> list[DocumentChunk]:
        """
        获取文档的分块
        
        Args:
            document_id: 文档ID
            
        Returns:
            文档分块列表
        """
        async with AsyncSessionLocal() as session:
            result = await session.execute(
                select(DocumentChunk)
                .where(DocumentChunk.document_id == document_id)
                .order_by(DocumentChunk.chunk_index)
            )
            # 转为 list 来匹配返回签名 List[DocumentChunk]
            return list(result.scalars().all())

    async def search_similar_chunks(self, query: str, organization_id: str, limit: int = 10) -> list[dict[str, Any]]:
        """
        搜索相似的文档分块（待实现 - 需要向量嵌入）

        Args:
            query: 查询文本
            organization_id: 组织ID
            limit: 返回结果数量限制

        Returns:
            相似分块列表
        """
        # 这里将实现基于向量的相似度搜索
        # 暂时返回空列表，等向量嵌入服务完成后再实现
        return []

document_service = DocumentParser()

